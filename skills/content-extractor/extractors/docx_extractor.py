"""Extract content from DOCX files."""

from typing import Optional


class DOCXExtractor:
    """Extracts text content from .docx files."""

    def is_available(self) -> bool:
        """Check if python-docx is installed."""
        try:
            import docx
            return True
        except ImportError:
            return False

    def extract(self, path: str) -> Optional[str]:
        """
        Extract text from a DOCX file.

        Args:
            path: Path to the .docx file

        Returns:
            Extracted text content, or None if extraction fails
        """
        if not self.is_available():
            return None

        try:
            import docx
            doc = docx.Document(path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)

            return "\n\n".join(paragraphs)

        except Exception:
            return None

    def extract_full(self, path: str) -> dict:
        """
        Extract full content from DOCX including metadata.

        Returns:
            dict with keys: text, paragraphs, page_count (estimate), tables
        """
        if not self.is_available():
            return None

        try:
            import docx
            doc = docx.Document(path)
            paragraphs = []
            tables_text = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        table_rows.append(" | ".join(row_texts))
                if table_rows:
                    tables_text.append("\n".join(table_rows))

            # Estimate page count from paragraph count (avg ~500 chars/page)
            total_chars = sum(len(p) for p in paragraphs)
            page_count = max(1, total_chars // 500)

            return {
                "text": "\n\n".join(paragraphs),
                "paragraphs": paragraphs,
                "page_count": page_count,
                "tables": tables_text,
                "metadata": {
                    "paragraph_count": len(paragraphs),
                    "table_count": len(tables_text)
                }
            }

        except Exception:
            return None
